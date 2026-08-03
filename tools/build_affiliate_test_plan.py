from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OUT = Path("docs/testplan/ke-hoach-kiem-thu-affiliate-marketplace-mvp1.docx")
OUT.parent.mkdir(parents=True, exist_ok=True)

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
LIGHT = "F2F4F7"
PALE_BLUE = "E8EEF5"
RED = "9B1C1C"
GOLD = "7A5A00"
GRAY = "666666"
WHITE = "FFFFFF"
BLACK = "111111"

doc = Document()
sec = doc.sections[0]
sec.page_width = Inches(8.5)
sec.page_height = Inches(11)
sec.top_margin = Inches(1.0)
sec.bottom_margin = Inches(1.0)
sec.left_margin = Inches(1.0)
sec.right_margin = Inches(1.0)
sec.header_distance = Inches(0.492)
sec.footer_distance = Inches(0.492)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(10.5)
normal.font.color.rgb = RGBColor.from_string(BLACK)
normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
normal.paragraph_format.space_after = Pt(5)
normal.paragraph_format.line_spacing = 1.1

for name, size, color, before, after in [
    ("Heading 1", 16, BLUE, 14, 7),
    ("Heading 2", 13, BLUE, 10, 5),
    ("Heading 3", 11.5, DARK_BLUE, 7, 4),
]:
    st = styles[name]
    st.font.name = "Calibri"
    st.font.size = Pt(size)
    st.font.bold = True
    st.font.color.rgb = RGBColor.from_string(color)
    st._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    st._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    st.paragraph_format.space_before = Pt(before)
    st.paragraph_format.space_after = Pt(after)
    st.paragraph_format.keep_with_next = True

for list_style in ["List Bullet", "List Number"]:
    st = styles[list_style]
    st.font.name = "Calibri"
    st.font.size = Pt(10.5)
    st.paragraph_format.left_indent = Inches(0.5)
    st.paragraph_format.first_line_indent = Inches(-0.25)
    st.paragraph_format.space_after = Pt(4)
    st.paragraph_format.line_spacing = 1.1


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=110, bottom=90, end=110):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_widths(table, widths):
    scale = 6.5 / sum(widths)
    widths = [width * scale for width in widths]
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    total = int(sum(widths) * 1440)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = Inches(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(headers, rows, widths, font_size=8.5):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_repeat_table_header(table.rows[0])
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, PALE_BLUE)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(font_size)
        r.font.color.rgb = RGBColor.from_string(INK)
    for ridx, row_data in enumerate(rows):
        cells = table.add_row().cells
        for i, value in enumerate(row_data):
            if ridx % 2 == 1:
                set_cell_shading(cells[i], "FAFBFC")
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i == 0 and len(headers) > 2 else WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(value))
            r.font.size = Pt(font_size)
    set_table_widths(table, widths)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    return table


def add_bullets(items, numbered=False):
    style = "List Number" if numbered else "List Bullet"
    for item in items:
        p = doc.add_paragraph(style=style)
        p.add_run(item)


def add_callout(label, text, fill=LIGHT, color=INK):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    set_repeat_table_header(table.rows[0])
    set_cell_shading(cell, fill)
    set_cell_margins(cell, 130, 160, 130, 160)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(label + ": ")
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(color)
    p.add_run(text)
    set_table_widths(table, [6.75])
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_page_field(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


header = sec.header
hp = header.paragraphs[0]
hp.text = "KẾ HOẠCH KIỂM THỬ HỆ THỐNG AFFILIATE MARKETPLACE — MVP1"
hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
hp.runs[0].font.size = Pt(8.5)
hp.runs[0].font.color.rgb = RGBColor.from_string(GRAY)
footer = sec.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
fr = fp.add_run("Nội bộ — Trang ")
fr.font.size = Pt(8.5)
fr.font.color.rgb = RGBColor.from_string(GRAY)
add_page_field(fp)

# Trang bìa
doc.add_paragraph().paragraph_format.space_after = Pt(18)
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(4)
r = p.add_run("KẾ HOẠCH KIỂM THỬ")
r.bold = True
r.font.size = Pt(25)
r.font.color.rgb = RGBColor.from_string(INK)
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(18)
r = p.add_run("Nền tảng Affiliate Marketplace nhãn trắng — MVP1")
r.font.size = Pt(15)
r.font.color.rgb = RGBColor.from_string(DARK_BLUE)

add_table(
    ["Thuộc tính", "Giá trị"],
    [
        ["Mã tài liệu", "KHKT-AFF-MVP1-001"],
        ["Phiên bản", "1.0 — Dự thảo để rà soát"],
        ["Ngày lập", "30/07/2026"],
        ["Người lập", "QA Lead"],
        ["Phạm vi nguồn", "Toàn bộ tài liệu trong thư mục docs/requirements/docsaff_v1"],
        ["Mô hình áp dụng", "Kiểm thử dựa trên rủi ro; tham chiếu IEEE 829 và IEEE 29119"],
        ["Trạng thái requirement", "Chưa sẵn sàng hoàn toàn; có điều kiện chặn được nêu trong mục 6 và 10"],
    ],
    [1.55, 5.2],
    9.5,
)
add_callout(
    "Kết luận điều hành",
    "Có thể chuẩn bị thiết kế kiểm thử theo module, nhưng chưa nên bắt đầu kiểm thử toàn bộ luồng tài chính đầu-cuối cho đến khi chốt phạm vi MVP1, hợp đồng cancel/refund, quy tắc reversal/adjustment, permission tài chính và hợp đồng API.",
    fill="FFF4E5",
    color=GOLD,
)

doc.add_page_break()

doc.add_heading("1. Tổng quan dự án", level=1)
add_table(
    ["Nội dung", "Mô tả"],
    [
        ["Tên dự án", "Nền tảng Affiliate Marketplace nhãn trắng"],
        ["Tên tính năng", "Marketplace đa Tenant, quản lý ưu đãi và ghi nhận commission/revenue share"],
        ["Mô tả", "Nền tảng cho phép mỗi Tenant vận hành marketplace mang thương hiệu riêng; quản lý Brand, Category, Offer, visibility; theo dõi click/order; tính số tiền tạm tính/chính thức; reporting, settlement và audit."],
        ["Vai trò", "Admin/Ops, Finance, Tenant Admin, Tenant Marketing/Operator, Tenant Viewer, End User, Brand System, Platform/Scheduler."],
        ["Hệ thống phụ thuộc", "DNS/SSL, Brand API/webhook, dịch vụ xác thực, cơ sở dữ liệu, scheduler, reporting pipeline; Tenant Loyalty API được hoãn sau MVP1."],
    ],
    [1.55, 5.2],
    9,
)
doc.add_heading("1.1 Mục tiêu kiểm thử", level=2)
add_bullets([
    "Xác nhận các chức năng MVP1 đáp ứng BRD, SRS, use case, business rule và quyết định thay đổi đã được phê duyệt.",
    "Bảo đảm cách ly dữ liệu theo Tenant và kiểm soát quyền truy cập đối với dữ liệu tài chính, transaction và audit.",
    "Xác nhận luồng click → order → Pending → Confirmed/Cancelled, tính commission/revenue share và settlement có tính đúng đắn, truy vết và chống trùng.",
    "Đánh giá khả năng phục hồi khi retry, timeout, event sai thứ tự, dữ liệu thiếu/sai và thay đổi cấu hình đồng thời.",
    "Xác nhận trải nghiệm web nhãn trắng, đa ngôn ngữ, responsive và các trạng thái loading/empty/error.",
    "Cung cấp bằng chứng chất lượng và khuyến nghị phát hành dựa trên rủi ro còn lại.",
])

doc.add_heading("2. Phạm vi kiểm thử", level=1)
doc.add_heading("2.1 Trong phạm vi", level=2)
scope_rows = [
    ["1", "Quản lý Tenant", "Onboard, domain/DNS/SSL, branding, locale, user Tenant Portal, trạng thái và audit.", "Cao"],
    ["2", "Quản lý Category", "Danh sách, tạo/sửa, đa ngôn ngữ, ngưng hoạt động/xóa mềm, dependency.", "Trung bình"],
    ["3", "Quản lý Brand và Offer", "CRUD, trạng thái, category mapping, commission rule, hiệu lực, content đa ngôn ngữ.", "Cao"],
    ["4", "Assignment và visibility", "Admin assign pool; Tenant ẩn/hiện; rule AND; direct-link protection.", "Cao"],
    ["5", "Landing page", "Nhận diện Tenant, branding, catalog, search/filter/sort, card, earn display text, My Orders và redirect.", "Cao"],
    ["6", "Theo dõi click", "Tạo click_id, context, redirect, chống thao tác trùng, correlation và logging.", "Cao"],
    ["7", "Order và transaction", "Brand order success, xác thực, validation, idempotency, matching, exception, trạng thái.", "Rất cao"],
    ["8", "Commission và revenue share", "Rule precedence/fallback, số tạm tính/chính thức, pending_days, Affiliate keep.", "Rất cao"],
    ["9", "Cancel/refund và finalization", "Event cancel/refund, unmatched event, scheduler, Pending→Confirmed/Cancelled.", "Rất cao"],
    ["10", "Reporting và settlement", "Dashboard, list/detail/export, kỳ đối soát, adjustment và trạng thái settlement.", "Rất cao"],
    ["11", "Tenant Portal", "User/role, visibility, earn display, transaction, export và audit theo Tenant.", "Cao"],
    ["12", "Audit và quan sát hệ thống", "Audit event, sensitive fields, correlation, exception log và khả năng truy vết.", "Cao"],
]
add_table(["STT", "Nhóm chức năng", "Phạm vi chính", "Ưu tiên"], scope_rows, [0.42, 1.55, 4.05, 0.73], 8.2)

doc.add_heading("2.2 Ngoài phạm vi MVP1", level=2)
add_bullets([
    "Cashback B1, tính điểm và cộng điểm cho End User.",
    "Kết nối Tenant Loyalty API và Point Posting Log.",
    "Lịch sử order/cashback ba tab dành cho End User.",
    "Cước CPC; click tracking chỉ phục vụ tracking, reporting và correlation.",
    "Ứng dụng di động native, kiểm thử cài đặt và thông báo đẩy.",
    "Các chức năng được ghi rõ là phase sau hoặc chưa được phê duyệt.",
])
add_callout("Lưu ý phạm vi", "Chỉ kiểm thử theo baseline CR-01: earn display là text; không có cashback B1, point posting, Tenant Loyalty API hoặc Banner Management trong MVP1. Mọi tài liệu/màn hình khác baseline phải được BA/PO xác nhận trước khi kiểm thử.", fill="FDECEC", color=RED)

doc.add_heading("2.3 Phạm vi kiểm thử theo từng nhóm chức năng", level=2)
feature_sections = [
    ("2.3.1 Quản lý Tenant", ["CRUD và state transition Draft/Active/Inactive.", "Tính duy nhất Tenant code/domain; DNS/SSL readiness; branding và locale.", "User Tenant Portal, account status, role và data scope.", "Ngưng hoạt động khi còn order/settlement là câu hỏi cần chốt."]),
    ("2.3.2 Category", ["CRUD, tìm kiếm/lọc, localized content và fallback vi-VN.", "Một category mặc định active trên mỗi Brand; dependency và xóa mềm.", "Policy sửa category code sau khi có dependency còn mở."]),
    ("2.3.3 Brand, Offer và commission", ["Brand/Offer lifecycle; URL/effective period; category mapping.", "Offer commission override và fallback về Category; commission_value > 0.", "Pending days và điều kiện active.", "Fixed commission Brand→Affiliate cần BA/PO chốt."]),
    ("2.3.4 Assignment và visibility", ["Admin assignment active và Tenant visibility visible theo phép AND.", "Giá trị mặc định khi Tenant chưa thao tác; Brand/Offer chưa assign phải ẩn.", "Thay đổi visibility giữa click và conversion; unassign khi còn rule liên quan."]),
    ("2.3.5 Landing page và click", ["Tenant resolution, catalog, search/filter/sort, card/detail và My Orders cơ bản.", "Hai locale, fallback, responsive, loading/empty/error và accessibility.", "Click idempotency, warning khi thiếu member_ref, redirect và source_component.", "Navigation Offer card, redirect fallback và SEO còn mở; Banner Management thuộc Phase 2."]),
    ("2.3.6 Order, transaction và exception", ["Auth/schema, click/context matching, idempotency và duplicate handling.", "Tạo Pending và provisional fields; reject/exception không tạo order.", "Exception type, list/detail, scope, resolution/reprocess cần hoàn thiện."]),
    ("2.3.7 Commission, cancel/refund và finalization", ["Category/Offer precedence, tenant_share_rate, Affiliate keep và rounding.", "Cancel/refund trước pending_until; scheduler finalization.", "Partial refund, refund sau Confirmed và reversal/adjustment là blocker."]),
    ("2.3.8 Reporting và settlement", ["Dashboard, filter, metric, data freshness, export và Tenant isolation.", "Settlement cycle theo hợp đồng; create/detail/status/adjustment/export.", "Metric dictionary, workflow approval và post-settlement adjustment còn mở."]),
    ("2.3.9 Tenant Portal", ["Role, user, assigned Brand, Offer visibility, earn display và transaction.", "Tenant-scoped reporting/export/audit; Viewer financial access là blocker.", "Earn display effective datetime, locale và precedence."]),
    ("2.3.10 Audit và phi chức năng", ["Audit thay đổi cấu hình quan trọng, before/after, actor và correlation.", "Hiệu năng, bảo mật, khả dụng, compatibility, accessibility và observability.", "Retention, failed action log và target định lượng chưa được chốt."]),
]
for heading, bullets in feature_sections:
    doc.add_heading(heading, level=3)
    add_bullets(bullets)

doc.add_heading("3. Chiến lược kiểm thử", level=1)
strategy_rows = [
    ["Kiểm thử chức năng", "Đối chiếu luồng chính, thay thế, ngoại lệ, state transition và acceptance criteria theo từng module.", "Tất cả"],
    ["Kiểm thử giao diện/trải nghiệm", "Bố cục, label, trạng thái control, responsive, hai locale, fallback, loading/empty/error.", "Tenant CMS, Tenant Portal, Landing page"],
    ["Kiểm thử validation", "Required, format, length, boundary, precision, effective date, duplicate, special characters.", "UI, API, dữ liệu"],
    ["Kiểm thử tích hợp", "DNS/SSL, Brand event, scheduler, reporting pipeline, settlement và audit correlation.", "Đầu-cuối"],
    ["Kiểm thử API", "Auth, schema, status/error, idempotency, retry, timeout, event ordering và negative contract.", "Order, cancel/refund"],
    ["Kiểm thử dữ liệu", "Persistence, unique constraint, state/history, amount calculation, soft delete, Tenant isolation.", "Các entity nghiệp vụ"],
    ["Kiểm thử permission", "Role-action-data matrix; truy cập chéo Tenant; field-level financial visibility; export.", "CMS và Tenant Portal"],
    ["Kiểm thử tương thích", "Trình duyệt/độ phân giải được phê duyệt; responsive desktop/mobile web.", "Giao diện web"],
    ["Kiểm thử khói", "Luồng sống còn sau mỗi bản triển khai: đăng nhập, catalog, click, ingest order, transaction view.", "Mỗi bản dựng"],
    ["Kiểm thử xác nhận sửa lỗi", "Chạy tập trung vùng thay đổi và dependency trực tiếp sau bản sửa.", "Theo thay đổi"],
    ["Kiểm thử hồi quy", "Bộ P0/P1 và luồng liên module; ưu tiên state, money, permission, integration.", "Trước phát hành"],
    ["Kiểm thử thăm dò và tiêu cực", "Event sai thứ tự, race condition, cấu hình đổi giữa luồng, malformed payload, truy cập trái phép.", "Vùng rủi ro cao"],
]
add_table(["Loại kiểm thử", "Cách tiếp cận", "Áp dụng"], strategy_rows, [1.55, 4.05, 1.15], 8.3)

doc.add_heading("3.1 Ưu tiên dựa trên rủi ro", level=2)
add_table(
    ["Mức", "Tiêu chí", "Nhóm ưu tiên"],
    [
        ["P0", "Sai tiền, mất/corrupt dữ liệu, lộ dữ liệu Tenant, hệ thống không giao dịch được.", "Order, commission/share, cancel/refund, settlement, permission, idempotency."],
        ["P1", "Hỏng luồng nghiệp vụ chính hoặc ảnh hưởng nhiều Tenant nhưng có giải pháp tạm.", "Tenant/Brand/Offer, visibility, landing page/click, reporting/export, audit."],
        ["P2", "Ảnh hưởng cục bộ, UX hoặc trường hợp ít gặp.", "Localization nhỏ, visual consistency, filter phụ, nội dung trợ giúp."],
    ],
    [0.6, 3.2, 2.95],
    8.6,
)

doc.add_heading("3.2 Tự động hóa và kiểm thử thủ công", level=2)
add_bullets([
    "Ưu tiên tự động hóa hợp đồng API, idempotency, calculation, state transition, Tenant isolation và bộ hồi quy P0/P1 ổn định.",
    "Duy trì kiểm thử thủ công cho UX, accessibility, exploratory, nội dung đa ngôn ngữ, reconciliation review và requirement còn biến động.",
    "Không tự động hóa các luồng chưa chốt business rule; đánh dấu blocked và theo dõi trong open question register.",
])

doc.add_heading("4. Các cấp độ kiểm thử", level=1)
add_table(
    ["Cấp độ", "Trách nhiệm", "Mục tiêu/đầu ra"],
    [
        ["Kiểm thử đơn vị (tham chiếu)", "Phát triển", "Calculation function, validation, repository/service, scheduler và permission guard đạt độ bao phủ do nhóm phát triển cam kết."],
        ["Kiểm thử tích hợp", "QA + Phát triển", "API, cơ sở dữ liệu, scheduler, reporting, DNS/SSL và audit hoạt động đúng hợp đồng."],
        ["Kiểm thử hệ thống", "QA", "Xác nhận toàn bộ chức năng và phi chức năng trên môi trường gần sản xuất."],
        ["Kiểm thử chấp nhận người dùng", "BA/PO, Ops, Finance, đại diện Tenant", "Xác nhận nghiệp vụ, báo cáo, settlement và vận hành theo hợp đồng."],
        ["Kiểm thử hồi quy", "QA", "Bảo vệ luồng P0/P1 trước phát hành và sau thay đổi lớn."],
        ["Kiểm thử khói", "QA/Tự động", "Xác nhận bản dựng đủ ổn định để bắt đầu kiểm thử sâu."],
    ],
    [1.55, 1.35, 3.85],
    8.7,
)

doc.add_heading("5. Môi trường kiểm thử", level=1)
add_table(
    ["Hạng mục", "Yêu cầu", "Trạng thái thông tin"],
    [
        ["Môi trường", "Tích hợp, kiểm thử hệ thống, chấp nhận người dùng; tách dữ liệu và credential.", "Tên/URL chưa cung cấp"],
        ["Giao diện", "CMS Admin, Tenant Portal, Landing page nhãn trắng.", "Công nghệ/phiên bản chưa cung cấp"],
        ["Dịch vụ phía máy chủ", "Tenant, catalog, tracking, order, commission, reporting, settlement, audit.", "Công nghệ/phiên bản chưa cung cấp"],
        ["Cơ sở dữ liệu", "Dữ liệu Tenant/Brand/Offer/click/order/rule/settlement/audit có thể kiểm chứng.", "Loại/phiên bản chưa cung cấp"],
        ["Điểm cuối API", "Order success; cancel/refund; xác thực Brand; mock/stub khi cần.", "Contract/URL chưa chốt"],
        ["Trình duyệt", "Chrome, Edge, Safari, Firefox theo ma trận được PO/Engineering phê duyệt.", "Phiên bản chưa chốt"],
        ["Thiết bị", "Desktop và mobile web theo breakpoint được thiết kế.", "Danh sách chưa chốt"],
        ["Tài khoản", "Admin/Ops, Finance, Tenant Admin, Marketing/Operator, Viewer, End User; ít nhất hai Tenant.", "Phải chuẩn bị"],
        ["Dữ liệu", "Nhiều Brand/Category/Offer/rule/currency/timezone/state; dataset duplicate/refund/exception.", "Phải thiết kế"],
        ["Quan sát", "Log, correlation ID, audit, exception queue, metric và quyền truy vấn dữ liệu.", "Quyền truy cập chưa chốt"],
    ],
    [1.4, 3.95, 1.4],
    8.2,
)
doc.add_heading("5.1 Bộ dữ liệu tối thiểu", level=2)
add_bullets([
    "Hai Tenant active và một Tenant inactive; domain/branding/locale khác nhau để kiểm tra cách ly.",
    "Brand/Offer ở Draft/Active/Inactive, có/không assignment, visibility visible/hidden và effective period biên.",
    "Category default/non-default; Offer có commission override và fallback Category.",
    "Order hợp lệ, duplicate, invalid click, invalid visibility, missing mapping/rule và event cancel/refund unmatched.",
    "Dữ liệu Pending/Confirmed/Cancelled và settlement ở mọi trạng thái được phê duyệt.",
])

doc.add_heading("6. Tiêu chí bắt đầu kiểm thử", level=1)
entry = [
    ["E1", "Baseline MVP1 được BA/PO phê duyệt và đồng bộ BRD/SRS/wireframe/CR.", "Bắt buộc", "Chưa đạt"],
    ["E2", "Chốt API order/cancel/refund, partial refund và reversal sau Confirmed.", "Bắt buộc cho order đầu-cuối", "Chưa đạt"],
    ["E3", "Chốt calculation, rounding, Tenant share missing và settlement adjustment.", "Bắt buộc cho tài chính", "Chưa đạt"],
    ["E4", "Permission matrix và Viewer financial visibility được phê duyệt.", "Bắt buộc cho bảo mật", "Chưa đạt"],
    ["E5", "Bản dựng triển khai thành công; khói kỹ thuật đạt; không có lỗi chặn môi trường.", "Bắt buộc", "Chờ"],
    ["E6", "Môi trường, account, credential, test data, log/audit access sẵn sàng.", "Bắt buộc", "Chờ"],
    ["E7", "Test case P0/P1 được review; traceability và dữ liệu kiểm thử hoàn tất.", "Bắt buộc", "Chờ"],
    ["E8", "NFR critical có target đo được và công cụ tải/quan sát sẵn sàng.", "Cho kiểm thử phi chức năng", "Chưa đạt"],
]
add_table(["Mã", "Tiêu chí", "Áp dụng", "Hiện trạng"], entry, [0.48, 4.35, 1.15, 0.77], 8.2)

doc.add_heading("7. Tiêu chí kết thúc kiểm thử", level=1)
exit_rows = [
    ["X1", "100% test case P0 và P1 thuộc phạm vi phát hành được thực thi; P2 đạt tối thiểu 95%.", "Báo cáo thực thi"],
    ["X2", "Không còn lỗi Chặn hoặc Nghiêm trọng mở; lỗi Cao còn lại phải có đánh giá rủi ro và phê duyệt PO.", "Báo cáo lỗi/phê duyệt"],
    ["X3", "Tỷ lệ đạt P0/P1 tối thiểu 95%; mọi case blocked có quyết định xử lý rõ.", "Chỉ số thực thi"],
    ["X4", "Hồi quy hoàn tất; smoke bản ứng viên phát hành đạt 100%.", "Kết quả bộ hồi quy/khói"],
    ["X5", "Không phát hiện truy cập chéo Tenant hoặc rò rỉ trường tài chính trái quyền.", "Bằng chứng permission"],
    ["X6", "Kết quả calculation/settlement khớp bộ ví dụ chuẩn do Finance/PO phê duyệt.", "Biên bản đối chiếu"],
    ["X7", "Các target phi chức năng thuộc release scope đạt hoặc có chấp nhận rủi ro.", "Báo cáo phi chức năng"],
    ["X8", "Báo cáo tổng kết kiểm thử, rủi ro tồn dư và khuyến nghị phát hành được phê duyệt.", "Biên bản kết thúc"],
]
add_table(["Mã", "Tiêu chí", "Bằng chứng"], exit_rows, [0.48, 4.95, 1.32], 8.4)

doc.add_heading("8. Sản phẩm bàn giao kiểm thử", level=1)
add_bullets([
    "Kế hoạch kiểm thử đã phê duyệt.",
    "Bộ test case và ma trận truy vết requirement–test case.",
    "Thiết kế dữ liệu kiểm thử và bộ dữ liệu chuẩn cho calculation/settlement.",
    "Bộ kiểm thử tự động và kết quả chạy khi được áp dụng.",
    "Báo cáo lỗi với bằng chứng, mức độ nghiêm trọng và requirement liên quan.",
    "Báo cáo thực thi theo chu kỳ, báo cáo hồi quy và báo cáo khói.",
    "Báo cáo tổng kết kiểm thử, rủi ro tồn dư và khuyến nghị phát hành.",
])

doc.add_heading("9. Phân tích rủi ro", level=1)
risks = [
    ["R01", "Nghiệp vụ", "Scope MVP1 không đồng nhất giữa CR và wireframe/SRS GUI.", "Rất cao", "Cao", "Đồng bộ CR-01/15 toàn bộ tài liệu; khóa baseline trước thiết kế chi tiết."],
    ["R02", "Nghiệp vụ", "Cancel/refund, partial refund và reversal chưa chốt.", "Rất cao", "Cao", "Workshop Product/Finance/Integration; phê duyệt contract và state transition."],
    ["R03", "Nghiệp vụ", "Calculation/rounding/adjustment còn rule mở.", "Rất cao", "Cao", "Finance cung cấp calculation catalog và ví dụ vàng."],
    ["R04", "Bảo mật", "Permission catalog và quyền Viewer xem tài chính chưa chốt.", "Rất cao", "Trung bình", "Phê duyệt role-action-data matrix; kiểm thử chéo Tenant bắt buộc."],
    ["R05", "Kỹ thuật", "API auth/error/idempotency/retry contract thiếu.", "Cao", "Cao", "API spec versioned; contract test và fault injection."],
    ["R06", "Dữ liệu", "Late event/race condition làm sai state hoặc settlement.", "Rất cao", "Trung bình", "Thiết kế concurrency/event-order suite; audit/correlation bắt buộc."],
    ["R07", "Kỹ thuật", "NFR không có target đo được.", "Cao", "Cao", "Chốt SLA/SLO, tải, percentile và điều kiện đo."],
    ["R08", "Phụ thuộc", "Brand sandbox/webhook và DNS/SSL không ổn định.", "Cao", "Trung bình", "Mock/stub kiểm soát được; health check; lịch phối hợp tích hợp."],
    ["R09", "Môi trường", "Thiếu log, audit hoặc quyền DB làm chậm điều tra.", "Trung bình", "Trung bình", "Checklist observability và quyền truy cập trước test cycle."],
    ["R10", "Tiến độ", "Nhiều module và open question gây rework test case.", "Cao", "Cao", "Thiết kế theo risk slice; freeze P0 trước; quản lý thay đổi và impact analysis."],
    ["R11", "Báo cáo", "Metric/cut-off/freshness chưa khóa.", "Cao", "Trung bình", "Metric dictionary; dataset reconciliation chuẩn."],
    ["R12", "Tuân thủ", "Retention/masking/audit policy chưa đầy đủ.", "Cao", "Trung bình", "Security/Compliance phê duyệt data governance matrix."],
]
add_table(["Mã", "Loại", "Mô tả", "Tác động", "Khả năng", "Giảm thiểu"], risks, [0.42, 0.78, 2.35, 0.75, 0.78, 1.67], 7.4)

doc.add_heading("10. Giả định, phụ thuộc và câu hỏi mở", level=1)
doc.add_heading("10.1 Giả định kiểm thử có kiểm soát", level=2)
add_bullets([
    "Kế hoạch hiện dùng phạm vi MVP1 đã chốt trong CR-01: không cashback B1, không tính/cộng điểm và không gọi Tenant Loyalty API.",
    "Trạng thái order MVP1 là Pending, Confirmed và Cancelled; trường hợp invalid đi exception queue và không tạo order.",
    "Các giả định chỉ dùng để lập kế hoạch nguồn lực; không thay thế quyết định nghiệp vụ hoặc acceptance criteria.",
])
doc.add_heading("10.2 Phụ thuộc", level=2)
add_bullets([
    "BA/PO/Finance/Engineering khóa requirement, calculation, API và permission.",
    "Brand cung cấp sandbox, credential, payload mẫu và khả năng phát event theo kịch bản.",
    "DevOps cung cấp môi trường, DNS/SSL, scheduler control, dữ liệu reset và observability.",
    "Nhóm phát triển cung cấp unit test evidence, deployment note và danh sách thay đổi.",
    "Đại diện Tenant/Ops/Finance sẵn sàng cho UAT và reconciliation review.",
])
doc.add_heading("10.3 Câu hỏi cho BA/PO và chủ sở hữu", level=2)
questions = [
    ["OQ01", "Baseline MVP1 có loại hoàn toàn các nội dung cashback/estimated points/loyalty endpoint khỏi UI không?", "PO/BA", "Chặn scope"],
    ["OQ02", "Contract cancel/refund dùng endpoint/event nào; response, error và idempotency ra sao?", "Product/Engineering", "Chặn API"],
    ["OQ03", "MVP1 có partial refund không; refund sau Confirmed/settlement xử lý thế nào?", "Finance/PO", "Chặn tiền"],
    ["OQ04", "Thiếu Tenant share rule sẽ exception hay tạo transaction với share bằng 0?", "Finance/PO", "Chặn tiền"],
    ["OQ05", "Commission Fixed Brand→Affiliate có thuộc MVP1 không?", "Finance/PO", "Chặn test data"],
    ["OQ06", "Brand bị unassign thì revenue share/earn display rule có hiệu lực thế nào?", "PO/Ops", "Chặn lifecycle"],
    ["OQ07", "Viewer được xem/export trường tài chính nào?", "PO/Security", "Chặn permission"],
    ["OQ08", "Settlement states, approval, auto/manual creation và adjustment formula là gì?", "Finance/Ops", "Chặn settlement"],
    ["OQ09", "Chỉ VND hay multi-currency; precision/rounding chuẩn là gì?", "Finance", "Chặn calculation"],
    ["OQ10", "Offer card mở detail hay redirect; policy redirect fallback khi tracking lỗi là gì?", "PO/UX", "Chặn UI"],
    ["OQ11", "Target latency, throughput, availability, freshness và export limit?", "Engineering/PO", "Chặn NFR"],
    ["OQ12", "Retention/masking cho member_ref, event, audit và exception?", "Security/Compliance", "Chặn compliance"],
    ["OQ13", "Idempotency key precedence, retention và payload conflict behavior?", "Engineering/Product", "Chặn API"],
    ["OQ14", "Exception lifecycle, SLA, owner và reprocess behavior?", "Ops/Product", "Chặn vận hành"],
    ["OQ15", "Browser/device/breakpoint chính thức cho MVP1?", "PO/UX", "Chặn compatibility"],
]
add_table(["Mã", "Câu hỏi", "Chủ trì", "Tác động"], questions, [0.48, 4.35, 1.12, 0.8], 8.1)

doc.add_heading("11. Kế hoạch nguồn lực", level=1)
add_callout("Cơ sở ước lượng", "Ước lượng sơ bộ cho một vòng kiểm thử hệ thống đầy đủ sau khi requirement P0 được khóa. Không bao gồm thời gian chờ xử lý blocker, hiệu năng quy mô lớn hoặc nhiều vòng UAT.", fill=LIGHT)
resources = [
    ["QA Lead", "Lập kế hoạch, risk/traceability, review, điều phối UAT và báo cáo.", "1", "30–35 ngày công"],
    ["QA chức năng Web", "CMS, Tenant Portal, landing page, role và localization.", "2", "55–65 ngày công"],
    ["QA API/Tích hợp", "Order, cancel/refund, idempotency, exception và scheduler.", "1", "35–45 ngày công"],
    ["QA dữ liệu/tài chính", "Calculation, reporting, reconciliation và settlement.", "1", "30–40 ngày công"],
    ["QA tự động hóa", "API/UI smoke và regression P0/P1.", "1", "30–40 ngày công"],
    ["Hỗ trợ hiệu năng/bảo mật", "Tải, security review và test chuyên sâu theo target.", "0,5–1", "10–20 ngày công"],
]
add_table(["Vai trò", "Trách nhiệm", "Số lượng", "Ước lượng"], resources, [1.3, 3.35, 0.75, 1.35], 8.5)
add_bullets([
    "Tổng quy mô khuyến nghị: 5–6 QA toàn thời gian trong giai đoạn cao điểm, có thể chia sẻ một nguồn lực hiệu năng/bảo mật.",
    "Thời lượng dự kiến: 10–12 tuần từ requirement freeze đến release verification.",
    "Một vòng hồi quy đầy đủ: 8–12 ngày công thủ công; mục tiêu giảm còn 4–6 ngày công khi bộ tự động hóa P0/P1 ổn định.",
])

doc.add_heading("12. Lịch kiểm thử dự kiến", level=1)
schedule = [
    ["1", "Rà soát và khóa requirement", "Tuần 1–2", "BA/PO/QA/Finance", "Baseline, OQ/CR được chốt"],
    ["2", "Hoàn thiện kế hoạch kiểm thử", "Tuần 2", "QA Lead", "Kế hoạch được phê duyệt"],
    ["3", "Thiết kế test case và dữ liệu", "Tuần 2–5", "QA", "Bộ P0/P1/P2, traceability"],
    ["4", "Chuẩn bị môi trường/tự động hóa", "Tuần 3–5", "QA/DevOps/Dev", "Môi trường, stub, smoke"],
    ["5", "Kiểm thử vòng 1", "Tuần 6–8", "QA", "Kết quả chức năng/tích hợp"],
    ["6", "Sửa lỗi và xác nhận", "Tuần 7–9", "Dev/QA", "Lỗi P0/P1 được đóng"],
    ["7", "Hồi quy và phi chức năng", "Tuần 9–10", "QA", "Báo cáo hồi quy/NFR"],
    ["8", "Hỗ trợ UAT", "Tuần 10–11", "QA/BA/Ops/Finance", "Biên bản UAT"],
    ["9", "Xác minh phát hành", "Tuần 12", "QA/DevOps", "Smoke và khuyến nghị phát hành"],
]
add_table(["STT", "Mốc", "Thời gian", "Chủ trì", "Đầu ra"], schedule, [0.42, 1.75, 0.9, 1.25, 2.43], 8.2)
add_callout("Điều kiện lịch", "Lịch phải được cập nhật sau khi chốt open question P0/P1 và có kế hoạch triển khai thực tế. Mọi trễ ở API cancel/refund hoặc calculation/settlement sẽ dịch chuyển các mốc kiểm thử đầu-cuối.", fill="FFF4E5", color=GOLD)

doc.add_heading("13. Chỉ số kiểm thử", level=1)
metrics = [
    ["Tiến độ", "Tổng số test case; đã thiết kế/review/thực thi; đạt; không đạt; bị chặn; chưa chạy.", "Hằng ngày trong chu kỳ"],
    ["Bao phủ", "Bao phủ requirement, business rule, role, API endpoint, state transition và nhóm rủi ro.", "Mỗi mốc review"],
    ["Chất lượng", "Tỷ lệ đạt; tỷ lệ blocked; defect density theo module; reopen rate.", "Hằng ngày/tuần"],
    ["Lỗi", "Tổng lỗi; phân bố Chặn/Nghiêm trọng/Cao/Trung bình/Thấp; tuổi lỗi; xu hướng mở/đóng.", "Hằng ngày"],
    ["Hiệu quả", "Tỷ lệ phát hiện lỗi trước UAT; thời gian xác nhận sửa lỗi; tỷ lệ tự động hóa P0/P1.", "Mỗi vòng"],
    ["Rò rỉ lỗi", "Số lỗi lọt UAT/sản xuất chia tổng lỗi của cùng release; phân tích nguyên nhân.", "Sau UAT/phát hành"],
    ["Tài chính", "Số phép tính đối chiếu; sai lệch amount; reconciliation variance và unresolved exception.", "Mỗi chu kỳ dữ liệu"],
]
add_table(["Nhóm", "Chỉ số", "Tần suất"], metrics, [1.05, 4.7, 1.0], 8.5)
doc.add_heading("13.1 Ngưỡng chất lượng đề xuất", level=2)
add_bullets([
    "100% requirement P0/P1 có ít nhất một test case và có kết quả trước quyết định phát hành.",
    "100% test case smoke đạt trên bản ứng viên phát hành.",
    "Không có lỗi Chặn/Nghiêm trọng mở; không có lỗi truy cập chéo Tenant.",
    "Sai lệch calculation so với bộ dữ liệu vàng bằng 0 trong phạm vi rule đã phê duyệt.",
    "Ngưỡng hiệu năng/khả dụng chỉ được ghi chính thức sau khi Product/Engineering chốt target.",
])

doc.add_heading("14. Khuyến nghị", level=1)
add_bullets([
    "Đóng các blocker requirement theo thứ tự: scope MVP1 → cancel/refund/reversal → calculation/settlement → permission → API contract → NFR.",
    "Thiết lập một register duy nhất cho open question/decision với owner, hạn chốt, tài liệu ảnh hưởng và trạng thái áp dụng.",
    "Xây calculation catalog có ví dụ vàng cho Category fallback, Offer override, Tenant share, Affiliate keep, rounding, cancel và adjustment.",
    "Áp dụng contract test cho Brand API và bộ event-order/idempotency tự động ngay từ giai đoạn tích hợp.",
    "Tạo ít nhất hai Tenant trong mọi môi trường để kiểm thử cách ly dữ liệu và permission liên tục.",
    "Đưa correlation ID xuyên suốt click–order–commission–settlement–audit để giảm thời gian điều tra.",
    "Tách smoke P0 chạy mỗi bản dựng và regression P0/P1 chạy theo lịch; ưu tiên tự động hóa API/calculation trước UI.",
    "Không phát hành luồng tài chính nếu Finance chưa ký duyệt ví dụ tính tiền và reconciliation output.",
])

doc.add_heading("15. Phê duyệt kế hoạch", level=1)
add_table(
    ["Vai trò", "Họ tên", "Trạng thái", "Ngày", "Ghi chú"],
    [
        ["QA Lead", "TBD", "Chờ phê duyệt", "TBD", ""],
        ["Product Owner", "TBD", "Chờ phê duyệt", "TBD", ""],
        ["Business Analyst", "TBD", "Chờ phê duyệt", "TBD", ""],
        ["Engineering Lead", "TBD", "Chờ phê duyệt", "TBD", ""],
        ["Finance/Ops", "TBD", "Chờ phê duyệt", "TBD", ""],
        ["Security/Compliance", "TBD", "Chờ phê duyệt", "TBD", ""],
    ],
    [1.45, 1.25, 1.35, 0.85, 1.85],
    8.5,
)

doc.add_heading("Phụ lục A — Tài liệu nguồn", level=1)
add_bullets([
    "BRD, functional analysis, function list và use case list của Affiliate Marketplace Platform.",
    "SRS tổng và các SRS module: Tenant, Category, Brand/Offer, Landing page, Order/Transaction, Reporting/Reconciliation, Tenant Portal và Audit.",
    "Modeling, activity/sequence diagram, wireframe và mockup trong thư mục docs/requirements/docsaff_v1.",
    "Danh sách change request đồng bộ tài liệu cập nhật ngày 22/07/2026.",
    "Kết quả requirement review của bộ tài liệu docs/requirements/docsaff_v1, cập nhật ngày 30/07/2026.",
])

# Ngắt dòng bảng và kiểm soát widow/orphan
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.keep_together = True
                p.paragraph_format.widow_control = True

doc.core_properties.title = "Kế hoạch kiểm thử Affiliate Marketplace MVP1"
doc.core_properties.subject = "Kế hoạch kiểm thử hợp nhất cho toàn bộ requirement trong docs/requirements/docsaff_v1"
doc.core_properties.author = "QA Copilot"
doc.core_properties.keywords = "kiểm thử, QA, affiliate marketplace, MVP1"
doc.save(OUT)
print(OUT.resolve())
